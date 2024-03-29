import docx
from docx.shared import Pt
import sqlite3
import datetime
import time

try:
    connection = sqlite3.connect("timetable.db")
    cursor = connection.cursor()
    cursor.execute("Select * from Chris")
except:
    print('Database is missing!')


document = docx.Document('Ausbildungsnachweis.docx')
style = document.styles['Normal']
font = style.font
font.name = 'Frutiger 45 Light'
font.size = Pt(12)



mon_list = ['Montag1','Montag2','Montag3','Montag4','Montag5','Montag6']
mtime_c_list =['mTime1','mTime2','mTime3','mTime4','mTime5','mTime6']
tue_list = ['Dienstag1','Dienstag2','Dienstag3','Dienstag4','Dienstag5','Dienstag6']
tuetime_c_list = ['tueTime1','tueTime2','tueTime3','tueTime4','tueTime5','tueTime6']
#wed_list = ['Mittwoch1','Mittwoch2','Mittwoch3','Mittwoch4','Mittwoch5', 'Mittwoch6']
#thu_list = ['Donnerstag1', 'Donnerstag2', 'Donnerstag3','Donnerstag4', 'Donnerstag5','Donnerstag6']
#fri_list = ['Freitag1','Freitag2','Freitag3','Freitag4','Freitag5']




def wr_mon_task(task_list):
    counter_task = 0
    counter_day = 0
    for table in document.tables: #Goes through each table,row,cell and written paragraph in the document
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for x in mon_list:
                        try:
                            if mon_list[counter_day] == paragraph.text: #when entry of the list matches parapgraph
                                paragraph.style = document.styles['Normal'] #set the style for the written data
                                paragraph.text = paragraph.text.replace(mon_list[counter_day], task_list[counter_task]) #replace the dummy-text with the task entry
                                counter_task = counter_task + 1
                                counter_day = counter_day + 1
                        except Exception as IndexError:
                            break

def wr_mon_time(time_list):
    counter_task = 0
    counter_day = 0
    for table in document.tables: #Goes through each table,row,cell and written paragraph in the document
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for x in mtime_c_list:
                        #print('Thats from time_c_list: ' + time_c_list[counter_day] + str(int(time_list[counter_task] )))
                        try:
                            if mtime_c_list[counter_day] == paragraph.text: #when entry of the list matches parapgraph
                                paragraph.style = document.styles['Normal'] #set the style for the written data
                                paragraph.text = paragraph.text.replace(mtime_c_list[counter_day], str(int(time_list[counter_task]))) #replace the dummy-text with the task entry
                                counter_task = counter_task + 1
                                counter_day = counter_day + 1
                        except Exception as IndexError:
                            break


def gettime_mon(weeknr):
    list = []  #Creates List
    cursor.execute("Select dayofweek,task,time FROM Chris WHERE weeknumber =" + ' ' +str(weeknr)) #Takes day, noted task and time from the db
    rows = cursor.fetchall()
    for row in rows:
        if row[0] == 'Monday': #filters for all entries with monday
            list.append(row[2]) #2 um die Zeit aus der Liste zu entnehmen
    return list #returns list only with all tasks

def gettask_mon(weeknr):
    list = []  #Creates List
    cursor.execute("Select dayofweek,task,time FROM Chris WHERE weeknumber =" + ' ' +str(weeknr)) #Takes day, noted task and time from the db
    rows = cursor.fetchall()
    for row in rows:
        if row[0] == 'Monday': #filters for all entries with monday
            list.append(row[1])
    return list #returns list only with all tasks



#def wr_sheet(weeknr):


#weeknr = input('WeekNr: ')
weeknr = '20'


task_list = None
time_list = None

datalist = [gettask_mon(weeknr),gettime_mon(weeknr)]
print(datalist[0], datalist[1])
wr_mon_task(datalist[0])
wr_mon_time(datalist[1])
print('Save now!\n'*4)
document.save('NewList.docx')
#wr_sheet()
