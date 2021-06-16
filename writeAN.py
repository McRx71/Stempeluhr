import docx
from docx.shared import Pt
import sqlite3
import datetime
import time

try:
    connection = sqlite3.connect("timetable.db")
    cursor = connection.cursor()
    cursor.execute("Select * from Chris")
except:
    print('Database is missing!')

document = docx.Document('Ausbildungsnachweis.docx')
style = document.styles['Normal']
font = style.font
font.name = 'Frutiger 45 Light'
font.size = Pt(12)

mon_list = ['Montag1', 'Montag2', 'Montag3', 'Montag4', 'Montag5', 'Montag6']
mtime_c_list = ['mTime1', 'mTime2', 'mTime3', 'mTime4', 'mTime5', 'mTime6']

tue_list = ['Dienstag1', 'Dienstag2', 'Dienstag3', 'Dienstag4', 'Dienstag5', 'Dienstag6']
tuetime_c_list = ['tueTime1', 'tueTime2', 'tueTime3', 'tueTime4', 'tueTime5', 'tueTime6']

wed_list = ['Mittwoch1', 'Mittwoch2', 'Mittwoch3', 'Mittwoch4', 'Mittwoch5', 'Mittwoch6']
wedtime_c_list = ['wedTime1', 'wedTime2', 'wedTime3', 'wedTime4', 'wedTime5', 'wedTime6']


thu_list = ['Donnerstag1', 'Donnerstag2', 'Donnerstag3','Donnerstag4', 'Donnerstag5','Donnerstag6']
thutime_c_list = ['thuTime1', 'thuTime2', 'thuTime3', 'thuTime4', 'thuTime5', 'thuTime6']

fri_list = ['Freitag1','Freitag2','Freitag3','Freitag4','Freitag5']
fritime_c_list = ['friTime1', 'friTime2', 'friTime3', 'friTime4', 'friTime5', 'friTime6']

def wr_mon_task(task_list):
    counter_task = 0
    counter_day = 0
    for table in document.tables:  # Goes through each table,row,cell and written paragraph in the document
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for x in mon_list:
                        try:
                            if mon_list[counter_day] == paragraph.text:  # when entry of the list matches parapgraph
                                paragraph.style = document.styles['Normal']  # set the style for the written data
                                paragraph.text = paragraph.text.replace(mon_list[counter_day], task_list[counter_task])  # replace the dummy-text with the task entry
                                counter_task = counter_task + 1
                                counter_day = counter_day + 1
                        except Exception as IndexError:
                            break


def wr_mon_time(time_list):
    counter_task = 0
    counter_day = 0
    for table in document.tables:  # Goes through each table,row,cell and written paragraph in the document
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for x in mtime_c_list:
                        # print('Thats from time_c_list: ' + time_c_list[counter_day] + str(int(time_list[counter_task] )))
                        try:
                            if mtime_c_list[counter_day] == paragraph.text:  # when entry of the list matches parapgraph
                                paragraph.style = document.styles['Normal']  # set the style for the written data
                                paragraph.text = paragraph.text.replace(mtime_c_list[counter_day], str(int(
                                    time_list[counter_task])))  # replace the dummy-text with the task entry
                                counter_task = counter_task + 1
                                counter_day = counter_day + 1
                        except Exception as IndexError:
                            break


def gettime_mon(weeknr):
    list = []  # Creates List
    cursor.execute("Select dayofweek,task,time FROM Chris WHERE weeknumber =" + ' ' + str(
        weeknr))  # Takes day, noted task and time from the db
    rows = cursor.fetchall()
    for row in rows:
        if row[0] == 'Monday':  # filters for all entries with monday
            list.append(row[2])  # 2 um die Zeit aus der Liste zu entnehmen
    return list  # returns list only with all tasks


def gettask_mon(weeknr):
    list = []  # Creates List
    cursor.execute("Select dayofweek,task,time FROM Chris WHERE weeknumber =" + ' ' + str(
        weeknr))  # Takes day, noted task and time from the db
    rows = cursor.fetchall()
    for row in rows:
        if row[0] == 'Monday':  # filters for all entries with monday
            list.append(row[1])
    return list  # returns list only with all tasks


# ----------------------------------------------------------Tuesday-----------------------------------------------------------------------------------------------

def wr_tue_task(task_list):
    counter_task = 0
    counter_day = 0
    for table in document.tables:  # Goes through each table,row,cell and written paragraph in the document
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for x in tue_list:
                        try:
                            if tue_list[counter_day] == paragraph.text:  # when entry of the list matches parapgraph
                                paragraph.style = document.styles['Normal']  # set the style for the written data
                                paragraph.text = paragraph.text.replace(tue_list[counter_day], task_list[
                                    counter_task])  # replace the dummy-text with the task entry
                                counter_task = counter_task + 1
                                counter_day = counter_day + 1
                        except Exception as IndexError:
                            break


def wr_tue_time(time_list):
    counter_task = 0
    counter_day = 0
    for table in document.tables:  # Goes through each table,row,cell and written paragraph in the document
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for x in tuetime_c_list:
                        try:
                            if tuetime_c_list[
                                counter_day] == paragraph.text:  # when entry of the list matches parapgraph
                                paragraph.style = document.styles['Normal']  # set the style for the written data
                                paragraph.text = paragraph.text.replace(tuetime_c_list[counter_day], str(int(
                                    time_list[counter_task])))  # replace the dummy-text with the task entry
                                counter_task = counter_task + 1
                                counter_day = counter_day + 1
                        except Exception as IndexError:
                            break


def gettime_tue(weeknr):
    list = []  # Creates List
    cursor.execute("Select dayofweek,task,time FROM Chris WHERE weeknumber =" + ' ' + str(
        weeknr))  # Takes day, noted task and time from the db
    rows = cursor.fetchall()
    for row in rows:
        if row[0] == 'Tuesday':  # filters for all entries with monday
            list.append(row[2])  # 2 um die Zeit aus der Liste zu entnehmen
    return list  # returns list only with all tasks


def gettask_tue(weeknr):
    list = []  # Creates List
    cursor.execute("Select dayofweek,task,time FROM Chris WHERE weeknumber =" + ' ' + str(
        weeknr))  # Takes day, noted task and time from the db
    rows = cursor.fetchall()
    for row in rows:
        if row[0] == 'Tuesday':  # filters for all entries with monday
            list.append(row[1])
    return list  # returns list only with all tasks


# -------------------------------------------------Wednesday--------------------------------------------------------------------------------------------------------

def wr_wed_task(task_list):
    counter_task = 0
    counter_day = 0
    for table in document.tables:  # Goes through each table,row,cell and written paragraph in the document
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for x in wed_list:
                        try:
                            if wed_list[counter_day] == paragraph.text:  # when entry of the list matches parapgraph
                                paragraph.style = document.styles['Normal']  # set the style for the written data
                                paragraph.text = paragraph.text.replace(wed_list[counter_day], task_list[
                                    counter_task])  # replace the dummy-text with the task entry
                                counter_task = counter_task + 1
                                counter_day = counter_day + 1
                        except Exception as IndexError:
                            break


def wr_wed_time(time_list):
    counter_task = 0
    counter_day = 0
    for table in document.tables:  # Goes through each table,row,cell and written paragraph in the document
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for x in wedtime_c_list:
                        try:
                            if wedtime_c_list[
                                counter_day] == paragraph.text:  # when entry of the list matches parapgraph
                                paragraph.style = document.styles['Normal']  # set the style for the written data
                                paragraph.text = paragraph.text.replace(wedtime_c_list[counter_day], str(int(
                                    time_list[counter_task])))  # replace the dummy-text with the task entry
                                counter_task = counter_task + 1
                                counter_day = counter_day + 1
                        except Exception as IndexError:
                            break


def gettime_wed(weeknr):
    list = []  # Creates List
    cursor.execute("Select dayofweek,task,time FROM Chris WHERE weeknumber =" + ' ' + str(
        weeknr))  # Takes day, noted task and time from the db
    rows = cursor.fetchall()
    for row in rows:
        if row[0] == 'Wednesday':  # filters for all entries with monday
            list.append(row[2])  # 2 um die Zeit aus der Liste zu entnehmen
    return list  # returns list only with all tasks


def gettask_wed(weeknr):
    list = []  # Creates List
    cursor.execute("Select dayofweek,task,time FROM Chris WHERE weeknumber =" + ' ' + str(
        weeknr))  # Takes day, noted task and time from the db
    rows = cursor.fetchall()
    for row in rows:
        if row[0] == 'Wednesday':  # filters for all entries with monday
            list.append(row[1])
    return list  # returns list only with all tasks

# ----------------------------------Thursday-----------------------------------------
def wr_thu_task(task_list):
    counter_task = 0
    counter_day = 0
    for table in document.tables:  # Goes through each table,row,cell and written paragraph in the document
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for x in thu_list:
                        try:
                            if thu_list[counter_day] == paragraph.text:  # when entry of the list matches parapgraph
                                paragraph.style = document.styles['Normal']  # set the style for the written data
                                paragraph.text = paragraph.text.replace(thu_list[counter_day], task_list[
                                    counter_task])  # replace the dummy-text with the task entry
                                counter_task = counter_task + 1
                                counter_day = counter_day + 1
                        except Exception as IndexError:
                            break


def wr_thu_time(time_list):
    counter_task = 0
    counter_day = 0
    for table in document.tables:  # Goes through each table,row,cell and written paragraph in the document
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for x in thutime_c_list:
                        try:
                            if thutime_c_list[counter_day] == paragraph.text:  # when entry of the list matches parapgraph
                                paragraph.style = document.styles['Normal']  # set the style for the written data
                                paragraph.text = paragraph.text.replace(thutime_c_list[counter_day], str(int(time_list[counter_task])))  # replace the dummy-text with the task entry
                                counter_task = counter_task + 1
                                counter_day = counter_day + 1
                        except Exception as IndexError:
                            break


def gettime_thu(weeknr):
    list = []  # Creates List
    cursor.execute("Select dayofweek,task,time FROM Chris WHERE weeknumber =" + ' ' + str(
        weeknr))  # Takes day, noted task and time from the db
    rows = cursor.fetchall()
    for row in rows:
        if row[0] == 'Thursday':  # filters for all entries with monday
            list.append(row[2])  # 2 um die Zeit aus der Liste zu entnehmen
    return list  # returns list only with all tasks


def gettask_thu(weeknr):
    list = []  # Creates List
    cursor.execute("Select dayofweek,task,time FROM Chris WHERE weeknumber =" + ' ' + str(
        weeknr))  # Takes day, noted task and time from the db
    rows = cursor.fetchall()
    for row in rows:
        if row[0] == 'Thursday':  # filters for all entries with monday
            list.append(row[1])
    return list  # returns list only with all tasks

# ------------------------------------------------------Friday-----------------------------------
def wr_fri_task(task_list):
    counter_task = 0
    counter_day = 0
    for table in document.tables:  # Goes through each table,row,cell and written paragraph in the document
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for x in fri_list:
                        try:
                            if fri_list[counter_day] == paragraph.text:  # when entry of the list matches parapgraph
                                paragraph.style = document.styles['Normal']  # set the style for the written data
                                paragraph.text = paragraph.text.replace(fri_list[counter_day], task_list[counter_task])  # replace the dummy-text with the task entry
                                counter_task = counter_task + 1
                                counter_day = counter_day + 1
                        except Exception as IndexError:
                            break


def wr_fri_time(time_list):
    counter_task = 0
    counter_day = 0
    for table in document.tables:  # Goes through each table,row,cell and written paragraph in the document
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for x in fritime_c_list:
                        try:
                            if fritime_c_list[counter_day] == paragraph.text:  # when entry of the list matches parapgraph
                                paragraph.style = document.styles['Normal']  # set the style for the written data
                                paragraph.text = paragraph.text.replace(fritime_c_list[counter_day], str(int(time_list[counter_task])))  # replace the dummy-text with the task entry
                                counter_task = counter_task + 1
                                counter_day = counter_day + 1
                        except Exception as IndexError:
                            break


def gettime_fri(weeknr):
    list = []  # Creates List
    cursor.execute("Select dayofweek,task,time FROM Chris WHERE weeknumber =" + ' ' + str(
        weeknr))  # Takes day, noted task and time from the db
    rows = cursor.fetchall()
    for row in rows:
        if row[0] == 'Friday':  # filters for all entries with monday
            list.append(row[2])  # 2 um die Zeit aus der Liste zu entnehmen
    return list  # returns list only with all tasks


def gettask_fri(weeknr):
    list = []  # Creates List
    cursor.execute("Select dayofweek,task,time FROM Chris WHERE weeknumber =" + ' ' + str(
        weeknr))  # Takes day, noted task and time from the db
    rows = cursor.fetchall()
    for row in rows:
        if row[0] == 'Friday':  # filters for all entries with monday
            list.append(row[1])
    return list  # returns list only with all tasks



# ----------------------------Start------------------------------------------------
weeknr = '18'

#task_list = None
#time_list = None
# ----------------------------------------Monday--------------------------------------
datalist = [gettask_mon(weeknr), gettime_mon(weeknr)]
print(datalist[0], datalist[1])
wr_mon_task(datalist[0])
wr_mon_time(datalist[1])
# ---------------------Tuesday----------------------------------------------------------
datalist = [gettask_tue(weeknr), gettime_tue(weeknr)]
wr_tue_task(datalist[0])
wr_tue_time(datalist[1])
# ----------------Wednesday----------------------------
datalist = [gettask_wed(weeknr), gettime_wed(weeknr)]
wr_wed_task(datalist[0])
wr_wed_time(datalist[1])

# ----------------Thursday----------------------------
datalist = [gettask_thu(weeknr), gettime_thu(weeknr)]
wr_thu_task(datalist[0])
wr_thu_time(datalist[1])


# ----------------Friday----------------------------
datalist = [gettask_fri(weeknr), gettime_fri(weeknr)]
print(datalist[0], datalist[1])
wr_fri_task(datalist[0])
wr_fri_time(datalist[1])


document.save('Ausbildungsnachweis_neu.docx')

