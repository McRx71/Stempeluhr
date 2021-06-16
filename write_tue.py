import docx
from docx.shared import Pt
import sqlite3
import datetime
import time

try:
    connection = sqlite3.connect("timetable.db")
    cursor = connection.cursor()
    cursor.execute("Select * from Chris")
except:
    print('Database is missing!')


document = docx.Document('Ausbildungsnachweis.docx')
style = document.styles['Normal']
font = style.font
font.name = 'Frutiger 45 Light'
font.size = Pt(12)




tue_list = ['Dienstag1','Dienstag2','Dienstag3','Dienstag4','Dienstag5','Dienstag6']
tuetime_c_list = ['tueTime1','tueTime2','tueTime3','tueTime4','tueTime5','tueTime6']


def wr_tue_task(task_list):
    counter_task = 0
    counter_day = 0
    for table in document.tables: #Goes through each table,row,cell and written paragraph in the document
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for x in tue_list:
                        try:
                            if tue_list[counter_day] == paragraph.text: #when entry of the list matches parapgraph
                                paragraph.style = document.styles['Normal'] #set the style for the written data
                                paragraph.text = paragraph.text.replace(tue_list[counter_day], task_list[counter_task]) #replace the dummy-text with the task entry
                                counter_task = counter_task + 1
                                counter_day = counter_day + 1
                        except Exception as IndexError:
                            break

def wr_tue_time(time_list):
    counter_task = 0
    counter_day = 0
    for table in document.tables: #Goes through each table,row,cell and written paragraph in the document
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for x in tuetime_c_list:
                        try:
                            if tuetime_c_list[counter_day] == paragraph.text: #when entry of the list matches parapgraph
                                paragraph.style = document.styles['Normal'] #set the style for the written data
                                paragraph.text = paragraph.text.replace(tuetime_c_list[counter_day], str(int(time_list[counter_task]))) #replace the dummy-text with the task entry
                                counter_task = counter_task + 1
                                counter_day = counter_day + 1
                        except Exception as IndexError:
                            break


def gettime_tue(weeknr):
    list = []  #Creates List
    cursor.execute("Select dayofweek,task,time FROM Chris WHERE weeknumber =" + ' ' +str(weeknr)) #Takes day, noted task and time from the db
    rows = cursor.fetchall()
    for row in rows:
        if row[0] == 'Tuesday': #filters for all entries with monday
            list.append(row[2]) #2 um die Zeit aus der Liste zu entnehmen
    return list #returns list only with all tasks

def gettask_tue(weeknr):
    list = []  #Creates List
    cursor.execute("Select dayofweek,task,time FROM Chris WHERE weeknumber =" + ' ' +str(weeknr)) #Takes day, noted task and time from the db
    rows = cursor.fetchall()
    for row in rows:
        if row[0] == 'Tuesday': #filters for all entries with monday
            list.append(row[1])
    return list #returns list only with all tasks







datalist = [gettask_tue(weeknr),gettime_tue(weeknr)]
wr_tue_task(datalist[0])
wr_tue_time(datalist[1])

